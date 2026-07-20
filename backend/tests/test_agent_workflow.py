"""Tests for the constrained Career Copilot workflow and persistence."""

from __future__ import annotations

import uuid
import json
from collections.abc import Generator
from datetime import datetime, timedelta, timezone
from io import BytesIO
from pathlib import Path

import pytest
from fastapi.testclient import TestClient
from docx import Document as DocxDocument
from sqlalchemy import create_engine, func, select
from sqlalchemy.orm import Session, sessionmaker

from backend.app.agents.career_copilot.schemas import (
    CalculateMatchScoreInput,
    CalculateMatchScoreOutput,
    GenerateApplicationMaterialOutput,
    RetrieveCandidateEvidenceInput,
    RetrieveCandidateEvidenceOutput,
    SaveAnalysisResultOutput,
)
from backend.app.agents.career_copilot.tools import create_tool_registry
from backend.app.api.dependencies import get_embedding_provider
from backend.app.application.agent_service import AgentRunService, recover_stale_agent_runs
from backend.app.application.analysis_service import AnalysisService
from backend.app.application.retrieval_service import RetrievalService
from backend.app.core.config import get_settings
from backend.app.infrastructure.database.base import Base
from backend.app.infrastructure.database.models import AgentRun, AgentStep, Analysis, User
from backend.app.infrastructure.database.session import create_session_factory, get_db_session
from backend.app.infrastructure.llm.parser import ExtractedAnalysis, build_final_analysis
from backend.app.main import app


class EmptyEmbeddingProvider:
    dimension = 1024

    def embed_texts(self, texts: list[str]) -> list[list[float]]:
        return [[1.0] + [0.0] * 1023 for _ in texts]

    def embed_query(self, query: str) -> list[float]:
        return [1.0] + [0.0] * 1023


def extracted_result() -> ExtractedAnalysis:
    return ExtractedAnalysis.model_validate({
        "job_requirements": {"core_skills": ["Python", "Redis"]},
        "matched_skills": ["Python"],
        "missing_skills": ["Redis"],
        "reasoning": ["Python has profile evidence"],
        "confidence": 0.8,
    })


@pytest.fixture
def agent_database(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> Generator[tuple[TestClient, sessionmaker[Session]], None, None]:
    engine = create_engine(f"sqlite+pysqlite:///{(tmp_path / 'agent.db').as_posix()}")
    Base.metadata.create_all(engine)
    session_factory = create_session_factory(engine)

    def override_session() -> Generator[Session, None, None]:
        with session_factory() as session:
            yield session

    monkeypatch.setattr(AnalysisService, "extract_job", lambda *_args: extracted_result())
    monkeypatch.setenv("DOCUMENT_STORAGE_PATH", str(tmp_path / "uploads"))
    get_settings.cache_clear()
    app.dependency_overrides[get_db_session] = override_session
    app.dependency_overrides[get_embedding_provider] = lambda: EmptyEmbeddingProvider()
    try:
        with TestClient(app) as client:
            yield client, session_factory
    finally:
        app.dependency_overrides.clear()
        get_settings.cache_clear()
        engine.dispose()


def create_profile_and_job(client: TestClient, headers: dict[str, str] | None = None) -> dict:
    client.post(
        "/api/v1/profiles",
        headers=headers,
        json={"name": "Candidate", "skills": ["Python"]},
    )
    return client.post(
        "/api/v1/jobs",
        headers=headers,
        json={"title": "Backend Engineer", "description": "Python and Redis"},
    ).json()


def test_workflow_transitions_and_persists_run_steps(agent_database) -> None:
    client, session_factory = agent_database
    job = create_profile_and_job(client)

    started = client.post("/api/v1/agent/runs", json={"job_id": job["id"]})
    assert started.status_code == 202
    assert started.json()["status"] == "pending"

    response = client.get(f"/api/v1/agent/runs/{started.json()['run_id']}")
    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "completed"
    assert payload["current_step"] == "end"
    assert [step["step_name"] for step in payload["steps"]] == [
        "validate_input",
        "extract_job_requirements",
        "retrieve_candidate_evidence",
        "calculate_score",
        "generate_analysis",
        "save_result",
    ]
    assert all(step["status"] == "completed" for step in payload["steps"])
    assert payload["result"]["analysis"]["score"] == 50

    with session_factory() as session:
        assert session.scalar(select(func.count()).select_from(AgentRun)) == 1
        assert session.scalar(select(func.count()).select_from(AgentStep)) == 6
        assert session.scalar(select(func.count()).select_from(Analysis)) == 1


def test_agent_run_is_user_isolated(agent_database) -> None:
    client, _ = agent_database
    owner = {"X-User-Email": "owner@example.test"}
    job = create_profile_and_job(client, owner)
    run_id = client.post(
        "/api/v1/agent/runs", headers=owner, json={"job_id": job["id"]}
    ).json()["run_id"]

    assert client.get(
        f"/api/v1/agent/runs/{run_id}",
        headers={"X-User-Email": "other@example.test"},
    ).status_code == 404
    assert client.post(
        "/api/v1/agent/runs",
        headers={"X-User-Email": "other@example.test"},
        json={"job_id": job["id"]},
    ).status_code == 404


def test_failed_run_is_saved_and_a_new_run_can_recover(
    agent_database,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    client, session_factory = agent_database
    job = create_profile_and_job(client)
    attempts = 0

    def fail_once(*_args):
        nonlocal attempts
        attempts += 1
        if attempts == 1:
            raise RuntimeError("structured extraction failed")
        return extracted_result()

    monkeypatch.setattr(AnalysisService, "extract_job", fail_once)
    first_id = client.post("/api/v1/agent/runs", json={"job_id": job["id"]}).json()["run_id"]
    first = client.get(f"/api/v1/agent/runs/{first_id}").json()
    assert first["status"] == "failed"
    assert first["steps"][-1]["status"] == "failed"
    assert first["error_message"] == "structured extraction failed"

    second_id = client.post("/api/v1/agent/runs", json={"job_id": job["id"]}).json()["run_id"]
    second = client.get(f"/api/v1/agent/runs/{second_id}").json()
    assert second["status"] == "completed"
    with session_factory() as session:
        assert session.scalar(select(func.count()).select_from(AgentRun)) == 2


def test_duplicate_active_run_is_rejected(agent_database) -> None:
    client, session_factory = agent_database
    job = create_profile_and_job(client)
    with session_factory() as session:
        user = session.scalar(select(User).where(User.email == "demo@example.com"))
        assert user is not None
        session.add(AgentRun(
            user_id=user.id,
            job_id=uuid.UUID(job["id"]),
            status="running",
            current_step="extract_job_requirements",
            result_json={},
        ))
        session.commit()

    duplicate = client.post("/api/v1/agent/runs", json={"job_id": job["id"]})
    assert duplicate.status_code == 409
    assert duplicate.json() == {
        "detail": "an active agent run already exists for this job"
    }


def test_timeout_is_persisted_without_analysis(agent_database) -> None:
    client, session_factory = agent_database
    job = create_profile_and_job(client)
    ticks = iter((0.0, 0.0, 2.0))

    with session_factory() as session:
        service = AgentRunService(
            session,
            "demo@example.com",
            analyzer=AnalysisService(object()),  # provider is not used by the patched extractor
            retrieval_service=RetrievalService(
                session, "demo@example.com", EmptyEmbeddingProvider()
            ),
            timeout_seconds=1.0,
            clock=lambda: next(ticks),
        )
        started = service.create_run(uuid.UUID(job["id"]))
        assert started.status == "pending"
        service.execute(started.run_id)

    with session_factory() as session:
        run = session.get(AgentRun, started.run_id)
        assert run is not None
        assert run.status == "timeout"
        assert run.error_message == "agent run exceeded 1 seconds"
        assert run.steps[0].status == "failed"
        assert session.scalar(select(func.count()).select_from(Analysis)) == 0


def test_startup_recovery_fails_stale_running_run(agent_database) -> None:
    client, session_factory = agent_database
    job = create_profile_and_job(client)
    now = datetime.now(timezone.utc)
    with session_factory() as session:
        user = session.scalar(select(User).where(User.email == "demo@example.com"))
        assert user is not None
        run = AgentRun(
            user_id=user.id,
            job_id=uuid.UUID(job["id"]),
            status="running",
            current_step="retrieve_candidate_evidence",
            result_json={},
            updated_at=now - timedelta(minutes=10),
        )
        run.steps.append(AgentStep(
            step_name="retrieve_candidate_evidence",
            status="running",
            input_summary="structured requirements",
        ))
        session.add(run)
        session.commit()
        run_id = run.id

    with session_factory() as session:
        assert recover_stale_agent_runs(session, 60, now=now) == 1

    with session_factory() as session:
        recovered = session.get(AgentRun, run_id)
        assert recovered is not None
        assert recovered.status == "failed"
        assert recovered.error_message == "agent run was interrupted before completion"
        assert recovered.steps[0].status == "failed"


def _resume_bytes() -> bytes:
    document = DocxDocument()
    document.add_heading("Skills", level=1)
    document.add_paragraph("Python FastAPI production development")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_resume_rag_agent_analysis_timeline_integration(
    agent_database,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    client, session_factory = agent_database
    job = create_profile_and_job(client)
    uploaded = client.post(
        "/api/v1/documents/upload",
        files={"file": ("resume.docx", _resume_bytes())},
    )
    assert uploaded.status_code == 201

    def extract_with_retrieval(_self, _title, _description, candidate_profile):
        profile = json.loads(candidate_profile)
        has_evidence = bool(profile["retrieved_resume_evidence"])
        return ExtractedAnalysis.model_validate({
            "job_requirements": {"core_skills": ["Python"]},
            "matched_skills": ["Python"] if has_evidence else [],
            "unverified_skills": [] if has_evidence else ["Python"],
            "reasoning": ["Retrieved resume evidence supports Python"] if has_evidence else [],
            "confidence": 0.9 if has_evidence else 0.2,
        })

    monkeypatch.setattr(AnalysisService, "extract_job", extract_with_retrieval)
    started = client.post("/api/v1/agent/runs", json={"job_id": job["id"]})
    timeline = client.get(f"/api/v1/agent/runs/{started.json()['run_id']}").json()

    assert timeline["status"] == "completed"
    assert len(timeline["steps"]) == 6
    assert all(step["status"] == "completed" for step in timeline["steps"])
    assert timeline["result"]["analysis"]["score"] == 100
    assert timeline["result"]["evidence"][0]["content"] == (
        "Python FastAPI production development"
    )
    with session_factory() as session:
        analysis = session.scalar(select(Analysis))
        assert analysis is not None
        assert analysis.score == 100
        assert analysis.evidence_json == timeline["result"]["evidence"]


def test_tool_registry_exposes_only_allow_list_and_checks_schemas() -> None:
    final = build_final_analysis(extracted_result())
    calculated = CalculateMatchScoreOutput(
        score=final.score,
        score_breakdown=final.score_breakdown,
        matched_skills=final.matched_skills,
        partial_skills=final.partial_skills,
        missing_skills=final.missing_skills,
        unverified_skills=final.unverified_skills,
    )
    registry = create_tool_registry(
        retrieve_handler=lambda _value: RetrieveCandidateEvidenceOutput(),
        calculate_handler=lambda _value: calculated,
        generate_handler=lambda _value: GenerateApplicationMaterialOutput(analysis=final),
        save_handler=lambda _value: SaveAnalysisResultOutput(analysis_id=uuid.uuid4()),
    )
    assert registry.names == (
        "retrieve_candidate_evidence",
        "calculate_match_score",
        "generate_application_material",
        "save_analysis_result",
    )
    assert registry.invoke(
        "retrieve_candidate_evidence",
        RetrieveCandidateEvidenceInput(requirements=[]),
        RetrieveCandidateEvidenceOutput,
    ).evidence == []
    with pytest.raises(KeyError, match="not allowed"):
        registry.get("shell")
    with pytest.raises(TypeError, match="invalid input schema"):
        registry.invoke(
            "retrieve_candidate_evidence",
            CalculateMatchScoreInput(extracted_analysis=extracted_result()),
            RetrieveCandidateEvidenceOutput,
        )
