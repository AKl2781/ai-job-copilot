"""Resume document extraction, chunking, upload, and isolation tests."""

from __future__ import annotations

from collections.abc import Generator
from io import BytesIO
from pathlib import Path
import uuid

import pymupdf
import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy import create_engine, select
from sqlalchemy.orm import Session, sessionmaker

from backend.app.application.document_processing import (
    chunk_resume_text,
    extract_docx_text,
    extract_pdf_text,
)
from backend.app.core.config import get_settings
from backend.app.api.dependencies import get_embedding_provider
from backend.app.infrastructure.database.base import Base
from backend.app.infrastructure.database.models import Document, DocumentChunk, User
from backend.app.infrastructure.database.session import create_session_factory, get_db_session
from backend.app.infrastructure.embedding.provider import EmbeddingServiceError
from backend.app.main import app


class TestEmbeddingProvider:
    dimension = 1024

    def embed_texts(self, texts: list[str]) -> list[list[float]]:
        return [[1.0] + [0.0] * 1023 for _ in texts]

    def embed_query(self, query: str) -> list[float]:
        return [1.0] + [0.0] * 1023


class FailingEmbeddingProvider(TestEmbeddingProvider):
    def embed_texts(self, texts: list[str]) -> list[list[float]]:
        raise EmbeddingServiceError("test embedding failure")


def _pdf_bytes() -> bytes:
    pdf = pymupdf.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "SUMMARY")
    page.insert_text((72, 96), "Backend engineer building reliable APIs.")
    page.insert_text((72, 132), "EXPERIENCE")
    page.insert_text((72, 156), "Example Co - Senior Engineer")
    data = pdf.tobytes()
    pdf.close()
    return data


def _docx_bytes() -> bytes:
    document = DocxDocument()
    document.add_heading("Skills", level=1)
    document.add_paragraph("Python, FastAPI, PostgreSQL")
    document.add_heading("Education", level=1)
    document.add_paragraph("BSc Computer Science")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


@pytest.fixture
def document_api(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> Generator[tuple[TestClient, sessionmaker[Session], Path], None, None]:
    engine = create_engine(f"sqlite+pysqlite:///{(tmp_path / 'documents.db').as_posix()}")
    Base.metadata.create_all(engine)
    session_factory = create_session_factory(engine)
    storage_root = tmp_path / "uploads"
    monkeypatch.setenv("DOCUMENT_STORAGE_PATH", str(storage_root))
    get_settings.cache_clear()

    def override_session() -> Generator[Session, None, None]:
        with session_factory() as session:
            yield session

    app.dependency_overrides[get_db_session] = override_session
    app.dependency_overrides[get_embedding_provider] = lambda: TestEmbeddingProvider()
    try:
        with TestClient(app) as client:
            yield client, session_factory, storage_root
    finally:
        app.dependency_overrides.clear()
        get_settings.cache_clear()
        engine.dispose()


def test_pdf_and_docx_text_extraction() -> None:
    assert "Backend engineer" in extract_pdf_text(_pdf_bytes())
    assert "Python, FastAPI" in extract_docx_text(_docx_bytes())


def test_chunker_uses_sections_and_paragraphs() -> None:
    chunks = chunk_resume_text(
        "Summary\n\nBackend engineer.\n\nExperience\n\nExample Co\nSenior Engineer"
    )

    assert [(chunk.section, chunk.content) for chunk in chunks] == [
        ("Summary", "Backend engineer."),
        ("Experience", "Example Co Senior Engineer"),
    ]


@pytest.mark.parametrize(
    ("filename", "content", "expected_sections"),
    [
        ("resume.pdf", _pdf_bytes(), {"Summary", "Experience"}),
        ("resume.docx", _docx_bytes(), {"Skills", "Education"}),
    ],
    ids=["pdf", "docx"],
)
def test_upload_parses_and_persists_chunks(
    document_api, filename: str, content: bytes, expected_sections: set[str]
) -> None:
    client, session_factory, _ = document_api
    response = client.post(
        "/api/v1/documents/upload",
        headers={"X-User-Email": "owner@example.test"},
        files={"file": (filename, content)},
    )

    assert response.status_code == 201, response.text
    payload = response.json()
    assert payload["filename"] == filename
    assert payload["status"] == "ready"
    assert {chunk["section"] for chunk in payload["chunks"]} == expected_sections
    assert [chunk["chunk_index"] for chunk in payload["chunks"]] == list(
        range(len(payload["chunks"]))
    )
    assert Path(payload["storage_path"]).read_bytes() == content

    with session_factory() as session:
        stored = session.get(Document, uuid.UUID(payload["id"]))
        assert stored is not None
        assert stored.status == "ready"
        assert len(session.scalars(select(DocumentChunk)).all()) == len(payload["chunks"])
        assert all(
            len(chunk.embedding or []) == 1024
            for chunk in session.scalars(select(DocumentChunk))
        )


def test_uploads_are_owned_and_stored_per_user(document_api) -> None:
    client, session_factory, storage_root = document_api
    first = client.post(
        "/api/v1/documents/upload",
        headers={"X-User-Email": "first@example.test"},
        files={"file": ("resume.docx", _docx_bytes())},
    ).json()
    second = client.post(
        "/api/v1/documents/upload",
        headers={"X-User-Email": "second@example.test"},
        files={"file": ("resume.docx", _docx_bytes())},
    ).json()

    assert first["user_id"] != second["user_id"]
    assert first["storage_path"] != second["storage_path"]
    assert Path(first["storage_path"]).is_relative_to(storage_root)
    assert Path(second["storage_path"]).is_relative_to(storage_root)

    with session_factory() as session:
        users = {user.email: user.id for user in session.scalars(select(User))}
        documents = {document.user_id for document in session.scalars(select(Document))}
        assert documents == {users["first@example.test"], users["second@example.test"]}


def test_upload_rejects_unsupported_file_type(document_api) -> None:
    client, _, _ = document_api
    response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("resume.txt", b"plain text")},
    )

    assert response.status_code == 415


def test_document_list_detail_and_chunks_are_user_scoped(document_api) -> None:
    client, _, _ = document_api
    owner_headers = {"X-User-Email": "owner@example.test"}
    uploaded = client.post(
        "/api/v1/documents/upload",
        headers=owner_headers,
        files={"file": ("resume.docx", _docx_bytes())},
    ).json()

    listed = client.get("/api/v1/documents", headers=owner_headers)
    assert listed.status_code == 200
    assert listed.json() == [
        {
            "id": uploaded["id"],
            "filename": "resume.docx",
            "file_type": "docx",
            "status": "ready",
            "chunk_count": 2,
            "created_at": uploaded["created_at"],
        }
    ]

    detail = client.get(f"/api/v1/documents/{uploaded['id']}", headers=owner_headers)
    assert detail.status_code == 200
    assert detail.json()["chunk_count"] == 2
    assert detail.json()["updated_at"] == uploaded["updated_at"]

    chunks = client.get(
        f"/api/v1/documents/{uploaded['id']}/chunks", headers=owner_headers
    )
    assert chunks.status_code == 200
    assert [chunk["section"] for chunk in chunks.json()] == ["Skills", "Education"]
    assert [chunk["chunk_index"] for chunk in chunks.json()] == [0, 1]
    assert set(chunks.json()[0]) == {"chunk_id", "section", "chunk_index", "content"}

    other_headers = {"X-User-Email": "other@example.test"}
    assert client.get("/api/v1/documents", headers=other_headers).json() == []
    assert client.get(
        f"/api/v1/documents/{uploaded['id']}", headers=other_headers
    ).status_code == 404
    assert client.get(
        f"/api/v1/documents/{uploaded['id']}/chunks", headers=other_headers
    ).status_code == 404


def test_duplicate_upload_returns_existing_document(document_api) -> None:
    client, session_factory, _ = document_api
    headers = {"X-User-Email": "owner@example.test"}
    content = _docx_bytes()
    first = client.post(
        "/api/v1/documents/upload",
        headers=headers,
        files={"file": ("resume.docx", content)},
    )
    duplicate = client.post(
        "/api/v1/documents/upload",
        headers=headers,
        files={"file": ("renamed.docx", content)},
    )

    assert first.status_code == 201
    assert duplicate.status_code == 200
    assert duplicate.json()["id"] == first.json()["id"]
    with session_factory() as session:
        assert len(session.scalars(select(Document)).all()) == 1


def test_upload_rejects_files_larger_than_ten_mb(document_api) -> None:
    client, session_factory, _ = document_api
    response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("large.pdf", b"x" * (10 * 1024 * 1024 + 1))},
    )

    assert response.status_code == 413
    with session_factory() as session:
        assert session.scalars(select(Document)).all() == []


def test_upload_sanitizes_path_components(document_api) -> None:
    client, _, storage_root = document_api
    response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("../../resume.docx", _docx_bytes())},
    )

    assert response.status_code == 201
    assert response.json()["filename"] == "resume.docx"
    assert Path(response.json()["storage_path"]).is_relative_to(storage_root)


def test_embedding_failure_marks_document_failed_without_saving_chunks(document_api) -> None:
    client, session_factory, _ = document_api
    app.dependency_overrides[get_embedding_provider] = lambda: FailingEmbeddingProvider()

    response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("resume.docx", _docx_bytes())},
    )

    assert response.status_code == 422
    with session_factory() as session:
        document = session.scalar(select(Document))
        assert document is not None
        assert document.status == "failed"
        assert session.scalars(select(DocumentChunk)).all() == []
