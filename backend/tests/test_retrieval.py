"""End-to-end semantic retrieval and user-isolation tests."""

from __future__ import annotations

from collections.abc import Generator
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import Session

from backend.app.api.dependencies import get_embedding_provider
from backend.app.core.config import get_settings
from backend.app.infrastructure.database.base import Base
from backend.app.infrastructure.database.session import create_session_factory, get_db_session
from backend.app.main import app


class SemanticTestProvider:
    dimension = 1024

    @staticmethod
    def _vector(text: str) -> list[float]:
        lowered = text.casefold()
        if "python" in lowered or "fastapi" in lowered:
            return [1.0, 0.0] + [0.0] * 1022
        return [0.0, 1.0] + [0.0] * 1022

    def embed_texts(self, texts: list[str]) -> list[list[float]]:
        return [self._vector(text) for text in texts]

    def embed_query(self, query: str) -> list[float]:
        return self._vector(query)


def _resume_bytes(skill: str) -> bytes:
    document = DocxDocument()
    document.add_heading("Skills", level=1)
    document.add_paragraph(skill)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


@pytest.fixture
def retrieval_api(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> Generator[TestClient, None, None]:
    engine = create_engine(f"sqlite+pysqlite:///{(tmp_path / 'retrieval.db').as_posix()}")
    Base.metadata.create_all(engine)
    session_factory = create_session_factory(engine)
    monkeypatch.setenv("DOCUMENT_STORAGE_PATH", str(tmp_path / "uploads"))
    get_settings.cache_clear()

    def override_session() -> Generator[Session, None, None]:
        with session_factory() as session:
            yield session

    app.dependency_overrides[get_db_session] = override_session
    app.dependency_overrides[get_embedding_provider] = lambda: SemanticTestProvider()
    try:
        with TestClient(app) as client:
            yield client
    finally:
        app.dependency_overrides.clear()
        get_settings.cache_clear()
        engine.dispose()


def test_retrieval_returns_top_chunks_by_cosine_similarity(retrieval_api) -> None:
    headers = {"X-User-Email": "owner@example.test"}
    python_document = retrieval_api.post(
        "/api/v1/documents/upload",
        headers=headers,
        files={"file": ("python.docx", _resume_bytes("Python and FastAPI backend"))},
    ).json()
    retrieval_api.post(
        "/api/v1/documents/upload",
        headers=headers,
        files={"file": ("design.docx", _resume_bytes("Product design and research"))},
    )

    response = retrieval_api.post(
        "/api/v1/retrieval/search",
        headers=headers,
        json={"query": "FastAPI backend experience", "top_k": 1},
    )

    assert response.status_code == 200
    assert response.json()[0]["document_id"] == python_document["id"]
    assert response.json()[0]["content"] == "Python and FastAPI backend"
    assert response.json()[0]["section"] == "Skills"
    assert response.json()[0]["score"] == pytest.approx(1.0)


def test_retrieval_never_returns_another_users_chunks(retrieval_api) -> None:
    owner_headers = {"X-User-Email": "owner@example.test"}
    other_headers = {"X-User-Email": "other@example.test"}
    owner_document = retrieval_api.post(
        "/api/v1/documents/upload",
        headers=owner_headers,
        files={"file": ("owner.docx", _resume_bytes("Python API development"))},
    ).json()
    retrieval_api.post(
        "/api/v1/documents/upload",
        headers=other_headers,
        files={"file": ("other.docx", _resume_bytes("Python FastAPI expert"))},
    )

    response = retrieval_api.post(
        "/api/v1/retrieval/search",
        headers=owner_headers,
        json={"query": "Python", "top_k": 10},
    )

    assert response.status_code == 200
    assert {match["document_id"] for match in response.json()} == {owner_document["id"]}
