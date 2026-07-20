"""Text extraction and structure-aware chunking for resume documents."""

from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO

import pymupdf
from docx import Document as DocxDocument


@dataclass(frozen=True)
class ResumeChunk:
    """A natural resume paragraph associated with its section."""

    content: str
    section: str


SECTION_NAMES = {
    "about": "About",
    "certifications": "Certifications",
    "certificates": "Certifications",
    "contact": "Contact",
    "education": "Education",
    "employment": "Experience",
    "experience": "Experience",
    "interests": "Interests",
    "languages": "Languages",
    "objective": "Objective",
    "professional experience": "Experience",
    "profile": "Profile",
    "projects": "Projects",
    "publications": "Publications",
    "qualifications": "Qualifications",
    "references": "References",
    "skills": "Skills",
    "summary": "Summary",
    "technical skills": "Skills",
    "work experience": "Experience",
    "个人简介": "个人简介",
    "个人信息": "个人信息",
    "专业技能": "专业技能",
    "工作经历": "工作经历",
    "工作经验": "工作经历",
    "教育经历": "教育经历",
    "教育背景": "教育经历",
    "技能": "技能",
    "求职意向": "求职意向",
    "项目经历": "项目经历",
    "项目经验": "项目经历",
}


def extract_pdf_text(data: bytes) -> str:
    """Extract reading-order text blocks from PDF bytes with PyMuPDF."""
    with pymupdf.open(stream=data, filetype="pdf") as pdf:
        pages: list[str] = []
        for page in pdf:
            blocks = page.get_text("blocks", sort=True)
            page_text = "\n\n".join(
                block[4].strip() for block in blocks if block[4].strip()
            )
            if page_text:
                pages.append(page_text)
    return "\n\n".join(pages)


def extract_docx_text(data: bytes) -> str:
    """Extract non-empty paragraphs and table cells from DOCX bytes."""
    document = DocxDocument(BytesIO(data))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(
                    paragraph.text.strip() for paragraph in cell.paragraphs
                )
    return "\n\n".join(text for text in paragraphs if text)


def _section_name(text: str) -> str | None:
    candidate = re.sub(r"^[\dIVXivx.()、\s]+", "", text.strip())
    candidate = candidate.rstrip(":：").strip()
    normalized = re.sub(r"\s+", " ", candidate).casefold()
    if normalized in SECTION_NAMES:
        return SECTION_NAMES[normalized]
    if text.rstrip().endswith((":", "：")) and 1 <= len(candidate) <= 80:
        return candidate
    return None


def _natural_paragraphs(text: str) -> list[str]:
    """Preserve source paragraphs while separating embedded section headings."""
    paragraphs: list[str] = []
    for block in re.split(r"\n\s*\n+", text.replace("\r\n", "\n")):
        lines = [re.sub(r"\s+", " ", line).strip() for line in block.splitlines()]
        lines = [line for line in lines if line]
        if not lines:
            continue
        current: list[str] = []
        for line in lines:
            if _section_name(line) is not None:
                if current:
                    paragraphs.append(" ".join(current))
                    current = []
                paragraphs.append(line)
            else:
                current.append(line)
        if current:
            paragraphs.append(" ".join(current))
    return paragraphs


def chunk_resume_text(text: str) -> list[ResumeChunk]:
    """Chunk a resume by recognized section headings and natural paragraphs."""
    chunks: list[ResumeChunk] = []
    section = "General"
    for paragraph in _natural_paragraphs(text):
        heading = _section_name(paragraph)
        if heading is not None:
            section = heading
            continue
        chunks.append(ResumeChunk(content=paragraph, section=section))
    return chunks
